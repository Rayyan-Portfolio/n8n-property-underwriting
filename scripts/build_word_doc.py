"""Generate Word documentation for the n8n underwriting pipeline."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING

OUT = r"d:\n8n\n8n-Property-Underwriting-Pipeline-Documentation.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    if bold:
        run.bold = True
    return para


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def code(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    para.paragraph_format.left_indent = Inches(0.25)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def build():
    doc = Document()
    set_doc_defaults(doc)

    # Title page
    title = doc.add_heading("Automated Property Underwriting Pipeline", level=0)
    title.alignment = 1
    sub = doc.add_paragraph("Complete Build & Operations Documentation")
    sub.alignment = 1
    p(doc, "Platform: n8n | Status: Demo-ready | June 2026", bold=False)
    doc.add_page_break()

    # 1 Overview
    h1(doc, "1. Overview")
    p(doc, "This system automates property lead underwriting from intake through Notion filing. A lead arrives via webhook POST or Google Form, is validated, enriched with free Census data, underwritten with deterministic formulas, given an AI-generated resale strategy, and saved as a deal page in Notion. The caller receives a JSON confirmation.")
    h2(doc, "What happens to each lead")
    numbered(doc, [
        "WF1 — Validates required fields and normalizes types.",
        "WF2 — Geocodes the address and pulls Census ACS neighborhood metrics.",
        "WF3 — Computes repair estimate, ARV, max offer, assignment offer, mortgage, deal score.",
        "WF4 — Generates structured resale strategy via Gemini 2.5 Flash (with template fallback).",
        "WF5 — Creates a Notion Deal Pipeline page with numbers in columns and strategy in the body.",
        "WF0 — Returns confirmation JSON with runId, address, dealScore, arvEstimate, arvConfidence.",
    ])
    h2(doc, "Architecture")
    code(doc, (
        "Google Form → Apps Script → ngrok → WF0 Webhook\n"
        "  → WF1 Validate → WF2 Enrich → WF3 Math → WF4 Strategy → WF5 Notion\n"
        "  → Build Response → Respond to Webhook\n"
        "\n"
        "Any system can also POST JSON directly to the webhook."
    ))
    h2(doc, "Workflows")
    table(doc, ["ID", "Name", "Role"], [
        ["WF0", "Orchestrator", "Public webhook, chains WF1–WF5, returns JSON"],
        ["WF1", "Validate & Normalize", "Quality gate, runId, type coercion"],
        ["WF2", "Enrich (Census)", "Geocode → FIPS → ACS metrics, 3 safe exit paths"],
        ["WF3", "Math Engine", "Deterministic underwriting formulas"],
        ["WF4", "AI Strategy", "Gemini structured JSON + template fallback"],
        ["WF5", "Notion Writer", "Database page + strategy body"],
    ])
    h2(doc, "Design principles")
    bullets(doc, [
        "Bottom-up build: each sub-workflow tested alone before wiring WF0.",
        "Math ≠ AI: dollar figures are auditable formulas; AI only handles language.",
        "Same output shape from every WF2 branch so downstream workflows need no extra branching.",
        "Fail safe: validation stops bad leads; enrichment and AI have fallbacks.",
    ])
    h2(doc, "Build order")
    p(doc, "WF1 → WF2 → WF3 → WF4 → WF5 → WF0")

    # 2 Setup
    h1(doc, "2. Environment Setup")
    h2(doc, "Prerequisites")
    table(doc, ["Requirement", "Notes"], [
        ["Docker", "n8n runs in a container"],
        ["n8n", "Editor at http://localhost:5678"],
        ["Census API key", "Free; ACS data (geocoder needs no key)"],
        ["Gemini API key", "Google AI Studio, free tier"],
        ["Notion account", "Internal integration + Deal Pipeline database"],
        ["ngrok", "Demo only — exposes local n8n to Google Apps Script"],
        ["Google account", "For Form + Apps Script (optional)"],
    ])
    h2(doc, "Run n8n in Docker")
    code(doc, (
        "docker run -d --name n8n \\\n"
        "  -p 5678:5678 \\\n"
        "  -v n8n_data:/home/node/.n8n \\\n"
        "  -e CENSUS_API_KEY=\"your-census-key\" \\\n"
        "  -e GEMINI_API_KEY=\"your-gemini-key\" \\\n"
        "  n8nio/n8n"
    ))
    p(doc, "Open http://localhost:5678 and complete first-time setup. Workflows and credentials persist in the n8n_data volume. Re-pass env vars if you recreate the container.")
    h2(doc, "Expose n8n with ngrok (Google Form demo)")
    code(doc, (
        "ngrok config add-authtoken YOUR_NGROK_TOKEN\n"
        "ngrok http 5678"
    ))
    p(doc, "Production webhook URL: https://YOUR-NGROK-HOST/webhook/YOUR-WEBHOOK-PATH")
    p(doc, "Update Apps Script WEBHOOK_URL whenever ngrok restarts.")
    h2(doc, "Publish workflows")
    p(doc, "Publish WF1–WF5 first, then WF0. Test URLs (/webhook-test/...) work while editing without publish.")

    # 3 API Keys
    h1(doc, "3. API Keys and Credentials")
    p(doc, "Never commit API keys to git. Store secrets in n8n credentials or environment variables.")
    h2(doc, "Census API key (free)")
    numbered(doc, [
        "Go to https://api.census.gov/data/key_signup.html",
        "Enter organization name and email.",
        "Check email and click the activation link.",
        "Save the key.",
    ])
    p(doc, "Used by WF2 ACS HTTP Request (key query parameter). Geocoder needs no key.")
    p(doc, "Gotcha: A bad Census key returns HTTP 200 + HTML. n8n treats this as success. WF2 Shape Enrichment detects HTML and falls back.")
    h2(doc, "Gemini API key (free tier)")
    numbered(doc, [
        "Go to https://aistudio.google.com",
        "Sign in → Get API key → Create API key.",
        "Copy the key.",
    ])
    code(doc, "POST https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash:generateContent?key=YOUR_KEY")
    p(doc, "Model: gemini-2.5-flash with responseMimeType application/json and fixed responseSchema.")
    h2(doc, "Notion integration token")
    numbered(doc, [
        "Go to https://www.notion.so/my-integrations",
        "New integration → Internal → name it (e.g. n8n Deal Pipeline).",
        "Copy the Internal Integration Secret (starts with ntn_).",
        "In n8n: Credentials → Notion API → save as Notion - Deal Pipeline.",
        "Connect integration to Deal Pipeline database (Section 4).",
    ])
    h2(doc, "Webhook secret (optional for demo)")
    code(doc, "[guid]::NewGuid().ToString(\"N\") + [guid]::NewGuid().ToString(\"N\")")
    p(doc, "Webhook node → Header Auth → header name x-webhook-secret. Callers send x-webhook-secret: YOUR_SECRET. Skipped for demo.")
    h2(doc, "Where keys live in n8n")
    table(doc, ["Key", "Storage"], [
        ["Notion", "n8n Notion credential"],
        ["Census", "Query param on ACS node or $env.CENSUS_API_KEY"],
        ["Gemini", "Query param on Gemini node or $env.GEMINI_API_KEY"],
        ["Webhook secret", "Header Auth credential on Webhook node"],
    ])

    # 4 Notion
    h1(doc, "4. Notion Deal Pipeline Database")
    h2(doc, "Create the database")
    numbered(doc, [
        "In Notion: new page → Table – Full page → name Deal Pipeline.",
        "Add properties with exact names (see table below).",
    ])
    table(doc, ["Property", "Type", "Notes"], [
        ["Address", "Title", "Default title column"],
        ["City/State", "Text", ""],
        ["Asking Price", "Number", ""],
        ["ARV Estimate", "Number", ""],
        ["ARV Confidence", "Select", "High, Medium, Low"],
        ["Repair Estimate", "Number", ""],
        ["Est. Monthly Payment", "Number", ""],
        ["Max Offer (MAO)", "Number", ""],
        ["Assignment Offer", "Number", ""],
        ["Deal Score", "Number", ""],
        ["Status", "Select", "New, Reviewed, Rejected"],
        ["Run ID", "Text", ""],
        ["Created", "Date", ""],
    ])
    h2(doc, "Connect the integration")
    p(doc, "Database → ⋯ menu → Connections → Connect to → select your integration. Without this, API returns object not found.")
    h2(doc, "Page body (WF5)")
    bullets(doc, [
        "Marketing Summary — strategy.marketingSummary",
        "Target Buyers — strategy.targetBuyerCriteria",
        "Asset Highlights — strategy.assetHighlights",
        "Negotiation Points — strategy.negotiationTalkingPoints",
    ])
    h2(doc, "WF5 property mapping")
    table(doc, ["Notion property", "n8n expression"], [
        ["Address (title)", "{{ $json.address }}"],
        ["City/State", "{{ $json.city }}, {{ $json.state }}"],
        ["Asking Price", "{{ $json.askingPrice }}"],
        ["ARV Estimate", "{{ $json.arvEstimate }}"],
        ["ARV Confidence", "{{ $json.arvConfidence }}"],
        ["Repair Estimate", "{{ $json.repairEstimate }}"],
        ["Est. Monthly Payment", "{{ $json.monthlyPI }}"],
        ["Max Offer (MAO)", "{{ $json.maxOffer }}"],
        ["Assignment Offer", "{{ $json.assignmentOffer }}"],
        ["Deal Score", "{{ $json.dealScore }}"],
        ["Status", "New"],
        ["Run ID", "{{ $json.runId }}"],
        ["Created", "{{ $json.receivedAt }}"],
    ])

    # 5 Data contract
    h1(doc, "5. Lead Data Contract")
    table(doc, ["Field", "Type", "Required", "Rules"], [
        ["address", "string", "yes", "Street address"],
        ["city", "string", "yes", ""],
        ["state", "string", "yes", "2-letter code (TX, not Texas)"],
        ["zip", "string", "yes", "5 digits"],
        ["beds", "number", "no", ""],
        ["baths", "number", "no", ""],
        ["sqft", "number", "yes", "Positive"],
        ["yearBuilt", "number", "yes", "1800–2100"],
        ["askingPrice", "number", "yes", "Positive; $250,000 coerced by WF1"],
        ["condition", "string", "no", "good / fair / poor"],
        ["notes", "string", "no", "Free text"],
    ])
    h2(doc, "Example payload")
    code(doc, (
        '{\n'
        '  "address": "120 Maple Ave",\n'
        '  "city": "Columbus",\n'
        '  "state": "OH",\n'
        '  "zip": "43004",\n'
        '  "beds": 3, "baths": 2, "sqft": 1500,\n'
        '  "yearBuilt": 2006, "askingPrice": 120000,\n'
        '  "condition": "good", "notes": "motivated seller"\n'
        '}'
    ))
    h2(doc, "Webhook confirmation response")
    code(doc, (
        '{\n'
        '  "status": "received",\n'
        '  "runId": "120-maple-ave-43004-...",\n'
        '  "address": "120 Maple Ave",\n'
        '  "dealScore": "0",\n'
        '  "arvEstimate": "225900",\n'
        '  "arvConfidence": "High"\n'
        '}'
    ))

    # 6 Workflows
    h1(doc, "6. Workflows — Build Guide")
    h2(doc, "WF1 — Validate & Normalize")
    p(doc, "Trigger: Execute Workflow Trigger (rename Lead In). Nodes: Lead In → Code Validate.")
    bullets(doc, [
        "Reads input.body ?? input (webhook wraps under body)",
        "Coerces numbers, trims text, uppercases state",
        "Rejects with Validation failed if invalid",
        "Adds runId, oneLineAddress, propertyAgeYears",
    ])
    h2(doc, "WF2 — Enrich (Census)")
    code(doc, (
        "Lead In → Geocode → Parse FIPS → IF (matched?)\n"
        "  TRUE  → ACS → Shape Enrichment\n"
        "  FALSE → Enrichment Fallback"
    ))
    bullets(doc, [
        "Geocode: geocoding.geo.census.gov (no key)",
        "ACS: api.census.gov/data/2023/acs/acs5 (key required)",
        "Shape Enrichment: detects bad-key HTML, filters -666666666 sentinels",
        "All three exit paths return identical field names",
    ])
    h2(doc, "WF3 — Math Engine")
    p(doc, "Trigger: Math In. Nodes: Math In → Code Underwrite.")
    table(doc, ["Metric", "Logic"], [
        ["Repair estimate", "$/sqft by condition + age bump × sqft"],
        ["ARV", "Average of regional median and asking × 1.15"],
        ["ARV confidence", "High / Medium / Low"],
        ["Max offer (MAO)", "ARV × 70% − repairs"],
        ["Assignment offer", "MAO − $10,000 fee"],
        ["Monthly P&I", "20% down, 7%, 30-year"],
        ["Deal score", "0–100 from MAO vs asking spread"],
    ])
    h2(doc, "WF4 — AI Strategy")
    p(doc, "Trigger: Math In. Nodes: Math In → HTTP Gemini → Code Parse Strategy.")
    bullets(doc, [
        "Gemini: structured JSON schema, On Error Continue",
        "Parse Strategy reads $('Math In') — trigger must be named Math In",
        "Fallback template if AI fails; arrays guarded with ?? []",
    ])
    h2(doc, "WF5 — Notion Writer")
    p(doc, "Trigger: Deal In. Nodes: Deal In → Notion Create Database Page. See Section 4 for mapping.")
    h2(doc, "WF0 — Orchestrator")
    code(doc, (
        "Webhook → Call WF1 → Call WF2 → Call WF3 → Call WF4 → Call WF5\n"
        "       → Build Response → Respond to Webhook"
    ))
    p(doc, "Build Response maps fields from Call WF4 (not WF5). Respond: First Incoming Item.")
    h2(doc, "Trigger naming checklist")
    table(doc, ["Workflow", "Trigger name"], [
        ["WF1", "Lead In (optional)"],
        ["WF2", "Lead In"],
        ["WF3", "Math In"],
        ["WF4", "Math In"],
        ["WF5", "Deal In"],
    ])

    # 7 Google Form
    h1(doc, "7. Google Form Demo Intake")
    table(doc, ["Form question", "JSON field"], [
        ["Address", "address"],
        ["City", "city"],
        ["State code in 2 letters (TX, not Texas).", "state"],
        ["Zip", "zip"],
        ["Beds / Baths / Sqft / Year Built", "beds, baths, sqft, yearBuilt"],
        ["Asking price", "askingPrice"],
        ["Condition", "condition"],
        ["Notes", "notes"],
    ])
    p(doc, "Extensions → Apps Script → onFormSubmit posts JSON to ngrok webhook URL.")
    p(doc, "Apps Script maps form titles to JSON keys. WF1 validates and normalizes.")
    p(doc, "Common mistake: question title mismatch leaves state empty. Use getState() with exact titles or /state/i fallback. Check Executions → Logs for Form titles found.")

    # 8 Testing
    h1(doc, "8. Testing")
    h2(doc, "PowerShell webhook test")
    code(doc, (
        '$body = @{ address="120 Maple Ave"; city="Columbus"; state="OH"; zip="43004"\n'
        '  beds=3; baths=2; sqft=1500; yearBuilt=2006; askingPrice=120000\n'
        '  condition="good" } | ConvertTo-Json\n'
        'Invoke-RestMethod -Uri "http://localhost:5678/webhook-test/lead-intake" `\n'
        '  -Method Post -Headers @{ "x-webhook-secret"="YOUR_SECRET" } `\n'
        '  -ContentType "application/json" -Body $body'
    ))
    h2(doc, "Edge-case matrix (all passed)")
    table(doc, ["#", "Case", "Expected"], [
        ["C1", 'state: "tx"', "Normalized to TX"],
        ["C2", 'state: "Texas"', "Rejected by WF1"],
        ["C3", 'askingPrice: "$250,000"', "Coerced to 250000"],
        ["C4", 'zip: "787"', "Rejected"],
        ["C5", "Missing sqft", "Rejected"],
        ["B1", "123 Main St, Austin, TX", "Low confidence, Notion created"],
        ["B2", "Google HQ, Mountain View", "Null median, Low/Medium ARV"],
        ["E1", "Invalid Gemini key", "Fallback template, chain completes"],
        ["A1", "4600 Silver Hill Rd, MD", "High enrichment, real Census data"],
    ])
    h2(doc, "Happy path checklist")
    numbered(doc, [
        "n8n execution — all nodes green",
        "Notion — new Deal Pipeline page",
        "Webhook — populated JSON response",
        "Apps Script log — 200 with JSON body (Form path)",
    ])

    # 9 Troubleshooting
    h1(doc, "9. Troubleshooting and Limitations")
    h2(doc, "Common issues")
    table(doc, ["Symptom", "Cause", "Fix"], [
        ["state must be 2-letter code", "Form title mismatch", "Fix Apps Script getState()"],
        ["askingPrice must be positive", "Form title mismatch", "Match Asking price title"],
        ["Referenced node doesn't exist", "Trigger not Math In", "Rename trigger in WF3/WF4"],
        ["Empty webhook response", "Respond reads WF5 output", "Build Response from Call WF4"],
        [".join is not a function", "Missing strategy arrays", "Guard arrays in Parse Strategy"],
        ["Notion object not found", "Integration not connected", "Connect integration to DB"],
        ["ngrok 404", "URL changed on restart", "Update WEBHOOK_URL"],
    ])
    h2(doc, "Known limitations")
    bullets(doc, [
        "Header Auth skipped for demo.",
        "No dedup — resubmit creates duplicate Notion pages.",
        "No error-handler workflow (WF9).",
        "n8n runs locally; ngrok required for Form intake.",
        "ARV based on Census regional median, not MLS comps.",
    ])
    h2(doc, "Build timeline")
    table(doc, ["Day", "Milestone", "Progress"], [
        ["Day 1", "Credentials, WF1, WF2, webhook skeleton", "45%"],
        ["Day 2", "WF3 Math, WF4 Gemini, WF5 Notion", "80%"],
        ["Day 3", "Full chain in WF0", "95%"],
        ["Day 4", "WF2 hardening, edge-case tests", "98%"],
        ["Day 5", "Bug fixes, Form intake, demo verified", "100% demo-ready"],
    ])

    h1(doc, "10. Result")
    p(doc, "Lead in → validated → enriched → underwritten → strategized → filed in Notion → JSON confirmation out. Project is demo-ready.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
